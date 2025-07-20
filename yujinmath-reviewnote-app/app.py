
import streamlit as st
import pandas as pd
from io import BytesIO
from zipfile import ZipFile
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

def set_korean_font(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    rFonts = style.element.rPr.rFonts
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

def generate_review_note(name, title, module1_dir, module2_dir, wrong1, wrong2):
    doc = Document()
    set_korean_font(doc)
    doc.add_paragraph(f"<{name}_{title}>").runs[0].bold = True

    if wrong1:
        doc.add_paragraph("<Module1>").runs[0].bold = True
        for q in wrong1:
            img_path = os.path.join(module1_dir, f"{q}.png")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=None)
                doc.add_paragraph("")

    if wrong2:
        doc.add_paragraph("<Module2>").runs[0].bold = True
        for q in wrong2:
            img_path = os.path.join(module2_dir, f"{q}.png")
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=None)
                doc.add_paragraph("")

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

st.title("📝 SAT 오답노트 생성기")

title = st.text_input("📌 문서 제목 (예: 25 SAT MATH S2 만점반 Mock3)")
module1_zip = st.file_uploader("📂 Module1 이미지 ZIP", type="zip", key="m1")
module2_zip = st.file_uploader("📂 Module2 이미지 ZIP", type="zip", key="m2")
excel_file = st.file_uploader("📋 오답노트 엑셀 업로드", type=["xlsx"])

if title and module1_zip and module2_zip and excel_file:
    with ZipFile(module1_zip) as m1zip:
        m1zip.extractall("/tmp/module1")
    with ZipFile(module2_zip) as m2zip:
        m2zip.extractall("/tmp/module2")

    df = pd.read_excel(excel_file)
    df.columns = df.columns.str.strip()  # 열 이름 공백 제거
    df = df.dropna(how='all')

    # 열 이름이 자동 감지되도록 처리
    col_map = {col.lower(): col for col in df.columns}
    name_col = col_map.get('이름') or col_map.get('name')
    mod1_col = col_map.get('module1')
    mod2_col = col_map.get('module2')

    file_buffer = BytesIO()
    with ZipFile(file_buffer, 'w') as zip_out:
        for _, row in df.iterrows():
            name = str(row[name_col]).strip()
            wrong1 = [] if pd.isna(row[mod1_col]) or row[mod1_col] == 'X' else str(row[mod1_col]).split(',')
            wrong2 = [] if pd.isna(row[mod2_col]) or row[mod2_col] == 'X' else str(row[mod2_col]).split(',')

            if not wrong1 and not wrong2:
                continue

            doc_stream = generate_review_note(name, title, "/tmp/module1", "/tmp/module2", wrong1, wrong2)
            zip_out.writestr(f"{name}_{title}.docx", doc_stream.read())

    file_buffer.seek(0)
    st.download_button("📥 오답노트 ZIP 다운로드", file_buffer, file_name=f"{title}_오답노트.zip")
